import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import glob
import xml.etree.ElementTree as ET
import os

# Zone IDs
Z0 = "MSZ075"
Z1 = "MSZ047"
Z2 = "MSZ050"
Z3 = "ALZ051"
Z4 = "LAZ037"
Z5 = "MSZ077"
Z6 = "MSZ079"

# List of URLs to scrape
urls = [
    f"https://forecast.weather.gov/MapClick.php?zoneid={Z0}",
    f"https://forecast.weather.gov/MapClick.php?zoneid={Z1}",
    f"https://forecast.weather.gov/MapClick.php?zoneid={Z2}",
    f"https://forecast.weather.gov/MapClick.php?zoneid={Z3}",
    f"https://forecast.weather.gov/MapClick.php?zoneid={Z4}",
    f"https://forecast.weather.gov/MapClick.php?zoneid={Z5}",
    f"https://forecast.weather.gov/MapClick.php?zoneid={Z6}"
]

# Loop through each URL
for i, url in enumerate(urls):
    # Send a GET request to the URL
    page = requests.get(url)

    # Parse the HTML content using Beautiful Soup
    soup = BeautifulSoup(page.content, "html.parser")

    # Find the div element with the ID "detailed-forecast-body"
    detailed_forecast = soup.find(id="detailed-forecast-body")

    # Find all the rows in the detailed forecast
    forecast_items = detailed_forecast.find_all(class_="row-forecast")

    # Create a new Word document
    doc = Document()

    # Set the document style to 'Normal'
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    # Add a title to the document
    title = doc.add_heading("Ag Zone Forecast", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add a subtitle to the document with the URL
    subtitle = doc.add_heading(f"NickelBlock Forecasting, {eval(f'Z{i}')}", level=1)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set the directory where the files will be saved
    directory = '/Users/nicke/Documents/'

    # Loop through each forecast item and add it to the document
    for item in forecast_items:
        # Get the time and description for this forecast item
        time = item.find(class_="forecast-label").text
        desc = item.find(class_="forecast-text").text

        # Create a new paragraph in the document
        p = doc.add_paragraph(style='Normal')

        # Add the time to the paragraph as bold text
        run = p.add_run(time)
        run.bold = True

        # Add the description to the paragraph as normal text
        p.add_run(": " + desc)

        # Set the paragraph alignment to left
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Save the forecast to a Word document
    if i == 0:
        filename = "forecast.docx"
    else:
        filename = f"forecast_{i}.docx"
    filepath = os.path.join(directory, filename)
    doc.save(filepath)

url = 'https://forecast.weather.gov/product.php?site=NWS&issuedby=JAN&product=AFD&format=CI&version=1&glossary=0'

response = requests.get(url)
soup = BeautifulSoup(response.content, 'html.parser')

pre_tag = soup.find('pre', class_='glossaryProduct')

if pre_tag:
    text = pre_tag.get_text()
    start_idx = text.find('DISCUSSION')
    end_idx = text.find('&&', start_idx)

    if start_idx != -1 and end_idx != -1:
        text = text[start_idx:end_idx].strip()
        # Replacing single line breaks with a space
        text = text.replace('\n', ' ')

        # Adding a new paragraph after every third period
        pattern = r"(.*?\..*?\..*?\.)"
        text = re.sub(pattern, r"\1\n\n", text)

        with open('/Users/nicke/Documents/WxOutput.txt', 'w') as f:
            f.write(text)
    else:
        print("Start or end index not found.")
else:
    print("Pre tag not found.")

# Get the content of the RSS feed
url = 'https://www.spc.noaa.gov/products/spcmdrss.xml'
response = requests.get(url)
xml_content = response.content

# Parse the XML content
root = ET.fromstring(xml_content)

# Find the first item in the RSS feed
item = root.find('.//item')

# Extract the title and description fields
title_element = item.find('title')
if title_element is not None:
    title = title_element.text
else:
    title = ""
description = item.find('description').text

# Write the title and description to an HTML file
with open('/Users/nicke/Documents/SPCoutput.html', 'w') as f:
    f.write(f'<h1>{title}</h1>\n')
    f.write(f'<p>{description}</p>\n')

# Set the URL of the webpage
url = 'https://www.spc.noaa.gov/products/md/'

# Send a GET request to the URL
response = requests.get(url)

# Check if the request was successful
if response.status_code == 200:
    # Parse the HTML content of the webpage using BeautifulSoup
    soup = BeautifulSoup(response.content, 'html.parser')

    # Find the link to the latest Mesoscale Discussion HTML file
    latest_md = soup.find('a', href=True, string='Latest Mesoscale Discussion')

    # Check if the link was found
    if latest_md:
        # Get the URL of the latest Mesoscale Discussion HTML file
        latest_md_url = latest_md['href']
        print(f'The URL of the latest Mesoscale Discussion is: {latest_md_url}')

        # Download the HTML file
        file_name = os.path.basename(latest_md_url)
        response = requests.get(latest_md_url)
        if response.status_code == 200:
            with open(file_name, 'wb') as f:
                f.write(response.content)
            print(f'Successfully downloaded {file_name}.')
        else:
            print(f'Request to {latest_md_url} failed with status code {response.status_code}.')
    else:
        print('Latest Mesoscale Discussion not found on the page.')
else:
    print(f'Request to {url} failed with status code {response.status_code}.')